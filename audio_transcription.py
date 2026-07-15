import os
import re
import sys
import json
import logging

# オフライン専用アプリのため、モデル読込時に一切ネットワークへ出ないよう強制する。
# faster_whisper / huggingface_hub のインポート前に設定する必要がある。
os.environ.setdefault("HF_HUB_OFFLINE", "1")
os.environ.setdefault("TRANSFORMERS_OFFLINE", "1")
# pyannote.audio は既定でテレメトリ（OTLP経由で otel.pyannote.ai へ使用状況を送信）が
# 有効になっている。オフライン専用アプリのため、pyannote.audio のインポート前に
# 明示的に無効化する（PYANNOTE_METRICS_ENABLED は pyannote.audio.telemetry.metrics が
# 参照する環境変数）。
os.environ.setdefault("PYANNOTE_METRICS_ENABLED", "false")

import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from faster_whisper import WhisperModel
from openpyxl import Workbook
from openpyxl.styles import Alignment, PatternFill
import datetime
import threading
import queue
import wave
import av
import docx
import winsound
import ctypes
import platform
import numpy as np
import sounddevice as sd
from faster_whisper.audio import decode_audio

# アプリケーション情報
APP_NAME = "TND_AudioTranscription"
APP_VERSION = "1.6.0"
APP_TITLE = f"TND audio_transcription v{APP_VERSION}"
APP_ICON_NAME = "TND_AudioTranscription01.ico"

# 低信頼区間とみなす avg_logprob のしきい値（これ未満はハイライト対象）
LOW_CONFIDENCE_LOGPROB = -0.8

# 単語登録キャッシュファイル名
HOTWORDS_FILE = "hotwords.json"
# 単語登録の上限数（hotwords枠223トークンに安全に収まる範囲）
MAX_HOTWORDS = 50

# アプリ設定ファイル名
SETTINGS_FILE = "settings.json"
DEFAULT_MODEL_NAME = "large-v3"

# 句読点を打たせるための呼び水プロンプト（句読点付きの文を与えることで
# モデルが句読点を出力しやすくなる。固有名詞は誤混入を避けるため含めない）
INITIAL_PROMPT = "こんにちは。本日は、よろしくお願いします。それでは、会議を始めます。"

# マイク録音の形式（16kHz / 16bit / モノラル固定）
RECORD_SAMPLE_RATE = 16000
RECORD_CHANNELS = 1
RECORD_DTYPE = "int16"

# 話者分離（同梱モデルがある場合のみ有効化）
DIARIZATION_MODEL_DIR = "models_diarization"
DIARIZATION_MODEL_REPO = "pyannote/speaker-diarization-community-1"
DIARIZATION_SAMPLE_RATE = 16000
DIARIZATION_UNKNOWN_SPEAKER = "不明"
DIARIZATION_NEAREST_THRESHOLD_SEC = 2.0

logger = logging.getLogger("audio_transcription")


class ProcessingCancelled(Exception):
    """ユーザー操作により処理がキャンセルされたことを示す例外"""
    pass


def cleanup_old_logs(log_dir, days=30):
    """30日より古いログファイルを削除する"""
    cutoff = datetime.datetime.now() - datetime.timedelta(days=days)
    for entry in os.listdir(log_dir):
        m = re.match(r"^app-(\d{8})\.log$", entry)
        if not m:
            continue
        try:
            file_date = datetime.datetime.strptime(m.group(1), "%Y%m%d")
        except ValueError:
            continue
        if file_date < cutoff:
            try:
                os.remove(os.path.join(log_dir, entry))
            except OSError:
                pass


def setup_logging():
    """ログ出力の初期設定（logs/app-YYYYMMDD.log に出力）"""
    app_dir = get_app_dir()
    log_dir = os.path.join(app_dir, "logs")
    os.makedirs(log_dir, exist_ok=True)
    try:
        cleanup_old_logs(log_dir)
    except OSError:
        pass
    log_file = os.path.join(log_dir, f"app-{datetime.datetime.now().strftime('%Y%m%d')}.log")

    handler = logging.FileHandler(log_file, encoding="utf-8")
    formatter = logging.Formatter(
        "%(asctime)s [%(levelname)s] %(name)s: %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S"
    )
    handler.setFormatter(formatter)

    logger.setLevel(logging.INFO)
    logger.addHandler(handler)

# ライセンス情報（アプリ内表示用）
LICENSE_TEXT = """\
─────────────────────────────────────
使用ライセンス情報

■ audio_transcription
MIT License
Copyright (c) 2024 HIROKI TANODA(TND)

本ソフトウェアおよび関連文書ファイル(以下「ソフトウェア」)のコピーを取得した
すべての人に対し、ソフトウェアを無制限に扱うことを無償で許可します。これには、
ソフトウェアのコピーを使用、複製、変更、結合、公開、頒布、サブライセンス、
および/または販売する権利、並びにソフトウェアを提供する相手に同じことを
許可する権利も無制限に含まれます。

上記の著作権表示および本許諾表示を、ソフトウェアのすべてのコピーまたは
重要な部分に記載するものとします。

THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, \
EXPRESS OR IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF \
MERCHANTABILITY, FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT.

■ Whisper Model (faster-whisper-large-v3)
MIT License  https://huggingface.co/Systran/faster-whisper-large-v3

■ faster-whisper
MIT License  Copyright (c) 2023 SYSTRAN
https://github.com/SYSTRAN/faster-whisper

■ OpenAI Whisper (Model)
MIT License  Copyright (c) 2022 OpenAI
https://github.com/openai/whisper

■ python-docx
MIT License  https://github.com/python-openxml/python-docx

■ lxml
BSD 3-Clause License  https://github.com/lxml/lxml

■ openpyxl
MIT License  Copyright (c) 2010 openpyxl
https://foss.heptapod.net/openpyxl/openpyxl

■ CTranslate2
MIT License  Copyright (c) 2019 OpenNMT
https://github.com/OpenNMT/CTranslate2

■ NumPy
BSD 3-Clause License  Copyright (c) 2005-2024, NumPy Developers
https://github.com/numpy/numpy

■ PyAV (FFmpeg Python bindings)
BSD 3-Clause License  Copyright (c) 2013, Mike Boers
https://github.com/PyAV-Org/PyAV

■ pyannote.audio（話者分離）
MIT License  Copyright (c) 2020 CNRS
https://github.com/pyannote/pyannote-audio

■ pyannote/speaker-diarization-community-1（話者分離モデル）
CC BY 4.0（帰属: pyannote）
https://huggingface.co/pyannote/speaker-diarization-community-1

■ PyTorch（話者分離の計算基盤・CPU版）
BSD 3-Clause License  Copyright (c) 2016- Facebook, Inc他
https://github.com/pytorch/pytorch

■ torchaudio
BSD 2-Clause License  Copyright (c) 2017 Facebook Inc.(Soumith Chintala)
https://github.com/pytorch/audio

■ python-sounddevice（マイク録音）
MIT License  Copyright (c) 2015-2025 Matthias Geier
https://github.com/spatialaudio/python-sounddevice

■ PortAudio（マイク録音の音声入出力ライブラリ）
MIT License  Copyright (c) 1999-2011 Ross Bencina and Phil Burk
https://www.portaudio.com/
─────────────────────────────────────\
"""

# 単語登録機能ヘルプテキスト
HOTWORDS_HELP_TEXT = """\
単語登録（カスタム辞書）機能について

■ 機能の概要
一般的な辞書には載っていない専門用語、社内用語、プロジェクト名、\
人名などをあらかじめ登録することで、AIによる文字起こしの誤変換を\
減らすことができます。

■ 登録数の目安と注意点
・登録上限： 最大50単語まで
・推奨登録数： 1回の会議につき 10〜20単語程度

⚠️ 重要：登録のコツ
単語を登録しすぎると、かえってAIが混乱し、関係のない会話まで\
無理やり登録単語に変換してしまう（誤認識が増える）可能性があります。
「どうしても間違えてほしくない重要な単語」に絞って登録するのが、\
最もきれいに文字起こしをするコツです。

■ 効果的な登録例
・固有名詞： 「TNDツール」「〇〇商事」「田野田」
・略語・業界用語： 「DX」「SaaS」「OJT」
・同音異義語の区別： 「あうとるっく（Outlook）」
  「きしょう（起床／気象）」など、文脈で間違えやすいもの\
"""


def get_hotwords_path():
    """単語登録ファイルのパスを取得"""
    app_dir = get_app_dir()
    return os.path.join(app_dir, HOTWORDS_FILE)


def load_hotwords():
    """保存済みの単語リストを読み込む"""
    path = get_hotwords_path()
    if os.path.exists(path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, list):
                return data
        except (json.JSONDecodeError, IOError):
            pass
    return []


def save_hotwords(words):
    """単語リストをJSONファイルに保存する"""
    path = get_hotwords_path()
    with open(path, "w", encoding="utf-8") as f:
        json.dump(words, f, ensure_ascii=False, indent=2)


def get_settings_path():
    """アプリ設定ファイルのパスを取得"""
    app_dir = get_app_dir()
    return os.path.join(app_dir, SETTINGS_FILE)


def load_settings():
    """保存済みのアプリ設定を読み込む"""
    path = get_settings_path()
    if os.path.exists(path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, dict):
                return data
        except (json.JSONDecodeError, IOError):
            pass
    return {}


def save_settings(settings):
    """アプリ設定をJSONファイルに保存する"""
    path = get_settings_path()
    with open(path, "w", encoding="utf-8") as f:
        json.dump(settings, f, ensure_ascii=False, indent=2)


def export_audio_segment(src_path, dst_path, start_sec, end_sec):
    """元音声から指定区間をストリームコピー（無劣化・再エンコードなし）で切り出す"""
    with av.open(src_path) as in_container:
        in_stream = in_container.streams.audio[0]
        with av.open(dst_path, mode="w") as out_container:
            out_stream = out_container.add_stream_from_template(in_stream)
            try:
                in_container.seek(int(start_sec / in_stream.time_base), stream=in_stream)
            except av.FFmpegError:
                pass
            offset = None
            for packet in in_container.demux(in_stream):
                if packet.pts is None:
                    continue
                t = float(packet.pts * in_stream.time_base)
                if t >= end_sec:
                    break
                if t < start_sec:
                    continue
                if offset is None:
                    offset = packet.pts
                packet.stream = out_stream
                packet.pts -= offset
                if packet.dts is not None:
                    packet.dts -= offset
                out_container.mux(packet)


def _spectral_gate(x, n_fft=512, hop=128):
    """1チャンクに対するスペクトルゲーティング（ノイズ抑制）"""
    win = np.hanning(n_fft)
    n_frames = 1 + (len(x) - n_fft) // hop
    idx = np.arange(n_fft)[None, :] + hop * np.arange(n_frames)[:, None]
    frames = x[idx] * win[None, :]
    spec = np.fft.rfft(frames, axis=1)
    mag = np.abs(spec).astype(np.float32)

    # 周波数ビンごとの雑音床を時間方向の下位パーセンタイルで推定（チャンクごとに適応）
    noise = np.percentile(mag, 15, axis=0)

    # ゲイン計算（過減算 1.5 倍、下限 0.1 で自然さを維持）
    gain = 1.0 - 1.5 * noise[None, :] / np.maximum(mag, 1e-10)
    gain = np.clip(gain, 0.1, 1.0)

    # 時間方向・周波数方向に軽く平滑化してミュージカルノイズを抑える
    g = gain
    g = (np.vstack([g[:1], g[:-1]]) + g + np.vstack([g[1:], g[-1:]])) / 3.0
    g = (np.hstack([g[:, :1], g[:, :-1]]) + g + np.hstack([g[:, 1:], g[:, -1:]])) / 3.0

    spec *= g

    rec = np.fft.irfft(spec, n=n_fft, axis=1) * win[None, :]
    y = np.zeros(len(x))
    norm = np.zeros(len(x))
    win_sq = win ** 2
    for i in range(n_frames):
        s = i * hop
        y[s:s + n_fft] += rec[i]
        norm[s:s + n_fft] += win_sq
    y[:n_frames * hop + n_fft] /= np.maximum(norm[:n_frames * hop + n_fft], 1e-8)
    # フレーム化で端数となった末尾はそのまま残す
    tail = n_frames * hop + n_fft - hop
    if tail < len(x):
        y[tail:] = x[tail:]
    return y


def preprocess_low_quality(audio, cancel_event=None, sampling_rate=16000):
    """雑音の多い音声を認識向けに前処理する（ノイズ抑制+音量正規化）。

    16kHz float32 モノラル波形を受け取り、60秒チャンクごとに
    スペクトルゲーティングを適用（メモリ使用を一定に保つ）し、
    最後に全体の音量を正規化した波形を返す。
    """
    x = np.asarray(audio, dtype=np.float64)
    if len(x) < 2048:
        return np.asarray(audio, dtype=np.float32)

    chunk = sampling_rate * 60
    y = np.empty(len(x))
    for s in range(0, len(x), chunk):
        if cancel_event is not None and cancel_event.is_set():
            raise ProcessingCancelled()
        seg = x[s:s + chunk]
        if len(seg) < 2048:
            y[s:s + len(seg)] = seg
        else:
            y[s:s + len(seg)] = _spectral_gate(seg)

    # RMS 正規化（目標 -20dBFS 相当、増幅は最大20倍まで）
    rms = np.sqrt(np.mean(y ** 2))
    if rms > 1e-8:
        y *= min(0.1 / rms, 20.0)
    y = np.clip(y, -1.0, 1.0)
    return y.astype(np.float32)


def detect_models():
    """modelsフォルダ内のHuggingFaceキャッシュ形式フォルダからモデル名を検出する"""
    model_dir = resource_path("models")
    if not os.path.isdir(model_dir):
        return []
    names = []
    for entry in sorted(os.listdir(model_dir)):
        full_path = os.path.join(model_dir, entry)
        if not os.path.isdir(full_path):
            continue
        m = re.match(r"^models--[^-]+(?:-[^-]+)*--faster-whisper-(.+)$", entry)
        if m:
            names.append(m.group(1))
    return names


def resolve_local_model_path(model_name):
    """指定モデルのローカルスナップショットフォルダのパスを返す（無効／不在なら None）。

    HuggingFace の名前解決やネットワークアクセスを一切介さず、同梱済みモデルの
    実ファイルがあるフォルダを直接特定する。model.bin の存在まで確認するため、
    不完全なインストールも None として検出できる。
    """
    model_dir = resource_path("models")
    if not os.path.isdir(model_dir):
        return None
    for entry in sorted(os.listdir(model_dir)):
        full_path = os.path.join(model_dir, entry)
        if not os.path.isdir(full_path):
            continue
        m = re.match(r"^models--[^-]+(?:-[^-]+)*--faster-whisper-(.+)$", entry)
        if not m or m.group(1) != model_name:
            continue
        snapshots = os.path.join(full_path, "snapshots")
        if not os.path.isdir(snapshots):
            return None
        candidates = []
        ref = os.path.join(full_path, "refs", "main")
        if os.path.isfile(ref):
            try:
                with open(ref, "r", encoding="utf-8") as f:
                    candidates.append(os.path.join(snapshots, f.read().strip()))
            except OSError:
                pass
        candidates.extend(
            os.path.join(snapshots, d)
            for d in sorted(os.listdir(snapshots))
            if os.path.isdir(os.path.join(snapshots, d))
        )
        for cand in candidates:
            if os.path.isdir(cand) and os.path.isfile(os.path.join(cand, "model.bin")):
                return cand
        return None
    return None


def resolve_diarization_model_path():
    """話者分離モデルの同梱キャッシュフォルダのパスを返す（無効／不在なら None）。

    Pipeline.from_pretrained(..., cache_dir=<この戻り値>) にそのまま渡せる、
    HuggingFace hub キャッシュのルート（models_diarization フォルダ自体）を返す。
    """
    cache_dir = resource_path(DIARIZATION_MODEL_DIR)
    if not os.path.isdir(cache_dir):
        return None
    model_dir = os.path.join(
        cache_dir, "models--" + DIARIZATION_MODEL_REPO.replace("/", "--")
    )
    snapshots = os.path.join(model_dir, "snapshots")
    if not os.path.isdir(snapshots):
        return None
    for entry in sorted(os.listdir(snapshots)):
        snap_path = os.path.join(snapshots, entry)
        if os.path.isdir(snap_path) and os.path.isfile(os.path.join(snap_path, "config.yaml")):
            return cache_dir
    return None


def run_selftest():
    """--selftest: GUIを起動せず、凍結EXEの同梱漏れを検出するための疎通確認を行う。

    1. faster-whisper: detect_models() + resolve_local_model_path() でモデル実体を解決できるか
    2. sounddevice: import + query_devices()（入力デバイス0件でも成功。import/DLLエラーのみ失敗）
    3. pyannote.audio: resolve_diarization_model_path() 解決 + Pipeline.from_pretrained() ロード
       + 5秒のダミー波形で pipeline を実行完走できるか（話者0人でも成功）

    結果は logs/selftest_YYYYMMDD_HHMMSS.log に書き出す。
    戻り値は終了コード（0=全成功 / 1=失敗あり）。
    """
    app_dir = get_app_dir()
    log_dir = os.path.join(app_dir, "logs")
    try:
        os.makedirs(log_dir, exist_ok=True)
    except OSError:
        pass
    log_path = os.path.join(
        log_dir, "selftest_{}.log".format(datetime.datetime.now().strftime("%Y%m%d_%H%M%S"))
    )

    results = []

    def record(name, ok, detail=""):
        results.append((name, ok, detail))

    # 1. faster-whisper: ローカルモデル実体の解決
    try:
        model_names = detect_models()
        if not model_names:
            raise RuntimeError("モデルが検出できません（detect_models() が空）")
        target_model = model_names[0]
        model_path = resolve_local_model_path(target_model)
        if not model_path:
            raise RuntimeError(f"モデル実体を解決できません: {target_model}")
        record("faster_whisper", True, f"model={target_model} path={model_path}")
    except Exception as e:
        record("faster_whisper", False, f"{type(e).__name__}: {e}")

    # 2. sounddevice: import + デバイス列挙（デバイス0件は許容、import/DLLエラーのみ失敗）
    try:
        devices = sd.query_devices()
        record("sounddevice", True, f"device_count={len(devices)}")
    except Exception as e:
        record("sounddevice", False, f"{type(e).__name__}: {e}")

    # 3. pyannote.audio: モデル解決 + ロード + ダミー波形での実行完走
    try:
        diarization_path = resolve_diarization_model_path()
        if not diarization_path:
            raise RuntimeError("話者分離モデルが検出できません")
        import torch
        from pyannote.audio import Pipeline
        pipeline = Pipeline.from_pretrained(DIARIZATION_MODEL_REPO, cache_dir=diarization_path)
        dummy_audio = (np.random.randn(DIARIZATION_SAMPLE_RATE * 5) * 0.01).astype(np.float32)
        waveform = torch.from_numpy(dummy_audio).unsqueeze(0)
        audio_input = {"waveform": waveform, "sample_rate": DIARIZATION_SAMPLE_RATE}
        diarization = pipeline(audio_input)
        annotation = getattr(diarization, "speaker_diarization", diarization)
        turn_count = sum(1 for _ in annotation.itertracks(yield_label=True))
        record("pyannote", True, f"path={diarization_path} turns={turn_count}")
    except Exception as e:
        record("pyannote", False, f"{type(e).__name__}: {e}")

    all_ok = all(ok for _, ok, _ in results)
    lines = [
        f"selftest {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"APP_VERSION={APP_VERSION}",
        "",
    ]
    for name, ok, detail in results:
        lines.append(f"[{'OK' if ok else 'NG'}] {name}: {detail}")
    lines.append("")
    lines.append("RESULT: {}".format("ALL_OK" if all_ok else "FAILED"))

    try:
        with open(log_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines) + "\n")
    except OSError:
        pass

    # --noconsole ビルドではコンソール未接続時に標準出力が使えないことがある
    # （ダブルクリック起動等）。ログファイルへの記録は上で完了済みのため、
    # コンソール出力の失敗で終了コードの返却を妨げないようにする。
    try:
        for line in lines:
            print(line)
    except Exception:
        pass

    return 0 if all_ok else 1


def decode_audio_pyav_16k_mono(path):
    """PyAVで音声を16kHzモノラルfloat32のnumpy配列にデコードする（話者分離の入力用）"""
    resampler = av.AudioResampler(format="fltp", layout="mono", rate=DIARIZATION_SAMPLE_RATE)
    chunks = []
    with av.open(path) as container:
        stream = container.streams.audio[0]
        for frame in container.decode(stream):
            for rf in resampler.resample(frame):
                chunks.append(rf.to_ndarray().reshape(-1))
        for rf in resampler.resample(None):
            chunks.append(rf.to_ndarray().reshape(-1))
    if not chunks:
        raise RuntimeError(f"音声フレームをデコードできませんでした: {path}")
    return np.concatenate(chunks).astype(np.float32)


def assign_speaker_to_segment(seg_start, seg_end, turns):
    """Whisperセグメントに話者を割り当てる（重なり合算最大→近傍2秒以内→不明）"""
    overlaps = {}
    for turn in turns:
        ov = min(seg_end, turn["end"]) - max(seg_start, turn["start"])
        if ov > 0:
            overlaps[turn["speaker"]] = overlaps.get(turn["speaker"], 0.0) + ov
    if overlaps:
        return max(overlaps.items(), key=lambda kv: kv[1])[0]

    center = (seg_start + seg_end) / 2.0
    best_speaker = None
    best_dist = None
    for turn in turns:
        dist = min(abs(center - turn["start"]), abs(center - turn["end"]))
        if best_dist is None or dist < best_dist:
            best_dist = dist
            best_speaker = turn["speaker"]
    if best_dist is not None and best_dist <= DIARIZATION_NEAREST_THRESHOLD_SEC:
        return best_speaker
    return DIARIZATION_UNKNOWN_SPEAKER


def get_app_dir():
    """アプリケーションのインストールディレクトリを取得"""
    if getattr(sys, 'frozen', False):
        # PyInstallerでビルドされた場合、EXEのあるフォルダ
        return os.path.dirname(sys.executable)
    else:
        # 開発時はスクリプトのあるフォルダ
        return os.path.dirname(os.path.abspath(__file__))


def resource_path(relative_path):
    """リソースファイルのパスを取得（モデル外部化対応）"""
    app_dir = get_app_dir()
    
    # EXEと同じフォルダを優先的に探す
    external_path = os.path.join(app_dir, relative_path)
    if os.path.exists(external_path):
        return external_path
    
    # PyInstallerの一時フォルダ（フォールバック）
    if hasattr(sys, '_MEIPASS'):
        meipass_path = os.path.join(sys._MEIPASS, relative_path)
        if os.path.exists(meipass_path):
            return meipass_path
    
    # 開発時のパス
    return os.path.join(os.path.abspath("."), relative_path)


class AudioTranscriptionApp:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title(APP_TITLE)
        self.root.geometry("500x940")
        self.root.resizable(True, True)

        self.input_file_paths = []
        self.output_folder_path = None
        self.is_processing = False
        self.cancel_event = threading.Event()
        self.hotwords_list = load_hotwords()
        self.settings = load_settings()

        # マイク録音用の状態
        self.is_recording = False
        self.record_stream = None
        self.record_queue = None
        self.record_writer_thread = None
        self.record_wave_file = None
        self.record_output_path = None
        self.record_start_dt = None
        self.record_level = 0
        self.record_error = None
        self.record_poll_id = None

        self.available_models = detect_models()
        self.selected_model_name = self.settings.get("model_name") or (
            self.available_models[0] if self.available_models else DEFAULT_MODEL_NAME
        )
        if self.available_models and self.selected_model_name not in self.available_models:
            self.selected_model_name = self.available_models[0]

        self.output_split_var = tk.BooleanVar(value=bool(self.settings.get("output_split", True)))
        self.output_txt_var = tk.BooleanVar(value=bool(self.settings.get("output_txt", False)))
        self.output_srt_var = tk.BooleanVar(value=bool(self.settings.get("output_srt", False)))
        self.output_docx_var = tk.BooleanVar(value=bool(self.settings.get("output_docx", False)))
        self.low_quality_var = tk.BooleanVar(value=bool(self.settings.get("low_quality_mode", False)))

        self.diarization_model_path = resolve_diarization_model_path()
        self.diarization_var = tk.BooleanVar(value=bool(self.settings.get("diarization", False)))

        # ウィンドウアイコンの設定
        self.set_window_icon()

        # モデルの存在確認
        self.check_model()

        self.setup_menu()
        self.setup_ui()
        self.populate_hotwords_listbox()

        self.root.protocol("WM_DELETE_WINDOW", self.on_close)

    def on_close(self):
        """ウィンドウを閉じる際の処理（処理中・録音中は確認する）"""
        if self.is_recording:
            if not messagebox.askyesno("確認", "録音中です。停止して終了しますか？"):
                return
            self.stop_recording(on_close=True)
        if self.is_processing:
            if messagebox.askyesno("確認", "処理が実行中です。中断して終了しますか？"):
                self.cancel_event.set()
                self.root.destroy()
        else:
            self.root.destroy()
    
    def set_window_icon(self):
        """ウィンドウアイコンを設定"""
        icon_path = os.path.join(get_app_dir(), APP_ICON_NAME)
        if os.path.exists(icon_path):
            try:
                self.root.iconbitmap(icon_path)
            except tk.TclError:
                pass  # アイコン読み込み失敗時は無視

    def check_model(self):
        """モデルの存在確認（フォルダの有無だけでなく、有効なモデル実体まで検証する）"""
        model_dir = resource_path("models")
        if not os.path.exists(model_dir):
            messagebox.showerror(
                "エラー",
                f"モデルフォルダが見つかりません。\n\n"
                f"期待されるパス:\n{model_dir}\n\n"
                f"アプリケーションを再インストールしてください。"
            )
            sys.exit(1)
        # 有効なモデル（model.bin まで揃ったスナップショット）が存在するか確認。
        # 不完全なインストールをここで検知し、実行時のネットワークダウンロード試行を防ぐ。
        if resolve_local_model_path(self.selected_model_name) is None:
            messagebox.showerror(
                "エラー",
                f"利用可能なモデルが見つかりません。\n\n"
                f"モデルフォルダ:\n{model_dir}\n\n"
                f"ZIPを展開せずにセットアップした場合や、モデルファイル "
                f"(約1.5〜3GB) のコピーが完了していない場合に発生します。\n"
                f"ZIPを完全に展開したうえで、アプリケーションを再インストールしてください。"
            )
            sys.exit(1)

    def setup_menu(self):
        """メニューバーを作成"""
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)

        # ヘルプメニュー
        help_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="ヘルプ", menu=help_menu)
        help_menu.add_command(label="単語登録機能について", command=self.show_hotwords_help)
        help_menu.add_separator()
        help_menu.add_command(label="サポート情報をコピー", command=self.copy_support_info)
        help_menu.add_separator()
        help_menu.add_command(label="ライセンス情報", command=self.show_license_info)

    def show_license_info(self):
        """ライセンス情報ダイアログを表示"""
        dialog = tk.Toplevel(self.root)
        dialog.title("ライセンス情報")
        dialog.geometry("560x450")
        dialog.resizable(True, True)
        dialog.transient(self.root)
        dialog.grab_set()

        # アイコンの設定
        icon_path = os.path.join(get_app_dir(), APP_ICON_NAME)
        if os.path.exists(icon_path):
            try:
                dialog.iconbitmap(icon_path)
            except tk.TclError:
                pass

        frame = ttk.Frame(dialog, padding="10")
        frame.pack(fill=tk.BOTH, expand=True)

        text_widget = tk.Text(frame, wrap=tk.WORD, font=("", 9))
        text_widget.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        scrollbar = ttk.Scrollbar(frame, orient=tk.VERTICAL, command=text_widget.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        text_widget.config(yscrollcommand=scrollbar.set)

        text_widget.insert(tk.END, LICENSE_TEXT)
        text_widget.config(state=tk.DISABLED)

        close_btn = ttk.Button(dialog, text="閉じる", command=dialog.destroy)
        close_btn.pack(pady=(0, 10))

    def show_hotwords_help(self):
        """単語登録機能ヘルプダイアログを表示"""
        dialog = tk.Toplevel(self.root)
        dialog.title("単語登録機能について")
        dialog.geometry("500x420")
        dialog.resizable(True, True)
        dialog.transient(self.root)
        dialog.grab_set()

        # アイコンの設定
        icon_path = os.path.join(get_app_dir(), APP_ICON_NAME)
        if os.path.exists(icon_path):
            try:
                dialog.iconbitmap(icon_path)
            except tk.TclError:
                pass

        frame = ttk.Frame(dialog, padding="10")
        frame.pack(fill=tk.BOTH, expand=True)

        text_widget = tk.Text(frame, wrap=tk.WORD, font=("", 10))
        text_widget.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        scrollbar = ttk.Scrollbar(frame, orient=tk.VERTICAL, command=text_widget.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        text_widget.config(yscrollcommand=scrollbar.set)

        text_widget.insert(tk.END, HOTWORDS_HELP_TEXT)
        text_widget.config(state=tk.DISABLED)

        close_btn = ttk.Button(dialog, text="閉じる", command=dialog.destroy)
        close_btn.pack(pady=(0, 10))

    def setup_ui(self):
        # メインフレーム
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 入力ファイル選択
        file_frame = ttk.LabelFrame(main_frame, text="入力ファイル", padding="10")
        file_frame.pack(fill=tk.X, pady=(0, 10))

        self.file_btn = ttk.Button(file_frame, text="選択", command=self.select_input_file)
        self.file_btn.pack(side=tk.RIGHT, padx=(10, 0))

        self.file_label = ttk.Label(file_frame, text="ファイルが選択されていません")
        self.file_label.pack(side=tk.LEFT, fill=tk.X, expand=True)

        # マイク録音
        record_frame = ttk.LabelFrame(main_frame, text="マイク録音", padding="10")
        record_frame.pack(fill=tk.X, pady=(0, 10))

        record_btn_row = ttk.Frame(record_frame)
        record_btn_row.pack(fill=tk.X, pady=(0, 5))

        self.record_start_btn = ttk.Button(record_btn_row, text="録音開始", command=self.start_recording)
        self.record_start_btn.pack(side=tk.LEFT, padx=(0, 5))

        self.record_stop_btn = ttk.Button(
            record_btn_row, text="停止", command=self.stop_recording, state=tk.DISABLED)
        self.record_stop_btn.pack(side=tk.LEFT, padx=(0, 5))

        self.record_time_label = ttk.Label(record_btn_row, text="00:00")
        self.record_time_label.pack(side=tk.LEFT, padx=(10, 0))

        self.record_level_bar = ttk.Progressbar(record_frame, mode='determinate', maximum=100)
        self.record_level_bar.pack(fill=tk.X)

        # モデル選択（検出されたモデルが2つ以上の場合のみ表示）
        if len(self.available_models) >= 2:
            model_frame = ttk.LabelFrame(main_frame, text="モデル", padding="10")
            model_frame.pack(fill=tk.X, pady=(0, 10))

            ttk.Label(model_frame, text="モデル: ").pack(side=tk.LEFT)

            self.model_combo = ttk.Combobox(
                model_frame, state="readonly", values=self.available_models
            )
            self.model_combo.set(self.selected_model_name)
            self.model_combo.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(5, 0))
            self.model_combo.bind("<<ComboboxSelected>>", self.on_model_selected)

        # 出力フォルダ選択
        folder_frame = ttk.LabelFrame(main_frame, text="出力フォルダ", padding="10")
        folder_frame.pack(fill=tk.X, pady=(0, 10))
        
        self.folder_btn = ttk.Button(folder_frame, text="選択", command=self.select_output_folder)
        self.folder_btn.pack(side=tk.RIGHT, padx=(10, 0))
        
        self.folder_label = ttk.Label(folder_frame, text="フォルダが選択されていません")
        self.folder_label.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        # プログレスバーフレーム
        progress_frame = ttk.LabelFrame(main_frame, text="処理状況", padding="10")
        progress_frame.pack(fill=tk.X, pady=(0, 10))
        
        self.status_label = ttk.Label(progress_frame, text="待機中...")
        self.status_label.pack(fill=tk.X, pady=(0, 5))
        
        self.progress_bar = ttk.Progressbar(progress_frame, mode='determinate', length=400)
        self.progress_bar.pack(fill=tk.X, pady=(0, 5))
        
        self.progress_detail = ttk.Label(progress_frame, text="")
        self.progress_detail.pack(fill=tk.X)
        
        # 単語登録フレーム
        hotwords_frame = ttk.LabelFrame(main_frame, text="単語登録（固有名詞・専門用語）", padding="10")
        hotwords_frame.pack(fill=tk.X, pady=(0, 10))

        # 入力行: テキスト入力 + 追加ボタン
        input_row = ttk.Frame(hotwords_frame)
        input_row.pack(fill=tk.X, pady=(0, 5))

        self.hotword_entry = ttk.Entry(input_row)
        self.hotword_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 5))
        self.hotword_entry.bind("<Return>", lambda e: self.add_hotword())

        self.add_btn = ttk.Button(input_row, text="追加", width=6, command=self.add_hotword)
        self.add_btn.pack(side=tk.LEFT, padx=(0, 5))

        self.remove_btn = ttk.Button(input_row, text="削除", width=6, command=self.remove_hotword)
        self.remove_btn.pack(side=tk.LEFT)

        # 登録数カウンター
        self.hotwords_count_label = ttk.Label(hotwords_frame, text=f"0 / {MAX_HOTWORDS} 件")
        self.hotwords_count_label.pack(anchor=tk.E, pady=(0, 3))

        # 登録済み単語リスト
        list_row = ttk.Frame(hotwords_frame)
        list_row.pack(fill=tk.X)

        self.hotwords_listbox = tk.Listbox(list_row, height=4, selectmode=tk.EXTENDED)
        self.hotwords_listbox.pack(side=tk.LEFT, fill=tk.X, expand=True)

        scrollbar = ttk.Scrollbar(list_row, orient=tk.VERTICAL, command=self.hotwords_listbox.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.hotwords_listbox.config(yscrollcommand=scrollbar.set)

        # 出力オプション
        output_options_frame = ttk.LabelFrame(main_frame, text="出力オプション", padding="10")
        output_options_frame.pack(fill=tk.X, pady=(0, 10))

        self.output_split_check = ttk.Checkbutton(
            output_options_frame, text="分割音声 (1分ごと・再生用) も出力",
            variable=self.output_split_var, command=self.on_output_option_changed
        )
        self.output_split_check.pack(anchor=tk.W)

        self.output_txt_check = ttk.Checkbutton(
            output_options_frame, text="テキスト (.txt) も出力",
            variable=self.output_txt_var, command=self.on_output_option_changed
        )
        self.output_txt_check.pack(anchor=tk.W)

        self.output_srt_check = ttk.Checkbutton(
            output_options_frame, text="字幕 (.srt) も出力",
            variable=self.output_srt_var, command=self.on_output_option_changed
        )
        self.output_srt_check.pack(anchor=tk.W)

        self.output_docx_check = ttk.Checkbutton(
            output_options_frame, text="Word (.docx) も出力",
            variable=self.output_docx_var, command=self.on_output_option_changed
        )
        self.output_docx_check.pack(anchor=tk.W)

        # 認識オプション
        recog_frame = ttk.LabelFrame(main_frame, text="認識オプション", padding="10")
        recog_frame.pack(fill=tk.X, pady=(0, 10))

        self.low_quality_check = ttk.Checkbutton(
            recog_frame, text="低品質音源モード（ノイズ抑制。雑音がひどい音源のみON推奨）",
            variable=self.low_quality_var, command=self.on_output_option_changed
        )
        self.low_quality_check.pack(anchor=tk.W)

        diarization_text = "話者分離を行う（Excel・Wordに話者を記載。処理時間が延びます）"
        if not self.diarization_model_path:
            diarization_text += "（モデル未導入）"
        self.diarization_check = ttk.Checkbutton(
            recog_frame, text=diarization_text,
            variable=self.diarization_var, command=self.on_output_option_changed
        )
        self.diarization_check.pack(anchor=tk.W)
        if not self.diarization_model_path:
            self.diarization_check.config(state=tk.DISABLED)

        # 実行ボタン・キャンセルボタン
        button_row = ttk.Frame(main_frame)
        button_row.pack(pady=10)

        self.run_btn = ttk.Button(button_row, text="文字起こし開始", command=self.start_processing)
        self.run_btn.pack(side=tk.LEFT, padx=(0, 5))

        self.cancel_btn = ttk.Button(button_row, text="キャンセル", command=self.cancel_processing, state=tk.DISABLED)
        self.cancel_btn.pack(side=tk.LEFT)

    def cancel_processing(self):
        """処理のキャンセルを要求"""
        self.cancel_event.set()
        self.cancel_btn.config(state=tk.DISABLED)
        self.status_label.config(text="キャンセル中...")

    @staticmethod
    def format_mmss(seconds):
        """秒数を MM:SS 形式の文字列に変換（録音経過時間表示用）"""
        sec = int(seconds)
        return f"{sec // 60:02d}:{sec % 60:02d}"

    def start_recording(self):
        """マイク録音を開始する"""
        if self.is_processing or self.is_recording:
            return
        if not self.output_folder_path:
            messagebox.showwarning("警告", "出力フォルダを選択してください。")
            return

        file_name = f"録音_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.wav"
        output_path = os.path.join(self.output_folder_path, file_name)

        try:
            os.makedirs(self.output_folder_path, exist_ok=True)
            wave_file = wave.open(output_path, 'wb')
            wave_file.setnchannels(RECORD_CHANNELS)
            wave_file.setsampwidth(2)
            wave_file.setframerate(RECORD_SAMPLE_RATE)
        except OSError:
            logger.exception(f"録音ファイルを作成できませんでした: {output_path}")
            messagebox.showerror(
                "エラー",
                "録音ファイルを作成できませんでした。\n出力フォルダの権限を確認してください。"
            )
            return

        self.record_queue = queue.Queue()
        self.record_error = None
        self.record_level = 0

        def on_audio_block(indata, frames, time_info, status):
            if status:
                logger.warning(f"録音ステータス: {status}")
            if len(indata):
                peak = int(np.abs(indata.astype(np.int32)).max())
                self.record_level = min(int(peak / 32767 * 100), 100)
            self.record_queue.put(bytes(indata))

        try:
            stream = sd.InputStream(
                samplerate=RECORD_SAMPLE_RATE, channels=RECORD_CHANNELS,
                dtype=RECORD_DTYPE, callback=on_audio_block
            )
            stream.start()
        except Exception:
            logger.exception("マイクのオープンに失敗しました。")
            wave_file.close()
            try:
                os.remove(output_path)
            except OSError:
                pass
            messagebox.showerror("エラー", "マイクを開けませんでした。マイクが接続されているか確認してください。")
            return

        def writer():
            while True:
                data = self.record_queue.get()
                if data is None:
                    break
                try:
                    wave_file.writeframes(data)
                except OSError:
                    logger.exception("録音データの書き込みに失敗しました。")
                    self.record_error = "録音データの書き込みに失敗しました。"
                    break

        self.record_stream = stream
        self.record_wave_file = wave_file
        self.record_output_path = output_path
        self.record_writer_thread = threading.Thread(target=writer, daemon=True)
        self.record_writer_thread.start()

        self.is_recording = True
        self.record_start_dt = datetime.datetime.now()
        self._set_recording_ui(True)
        logger.info(f"録音を開始しました: {output_path}")
        self._poll_recording()

    def _poll_recording(self):
        """録音中の経過時間・レベルメーターを定期更新する"""
        if not self.is_recording:
            return
        elapsed = (datetime.datetime.now() - self.record_start_dt).total_seconds()
        self.record_time_label.config(text=self.format_mmss(elapsed))
        self.record_level_bar['value'] = self.record_level
        if self.record_error:
            message = self.record_error
            self.record_error = None
            self._finalize_recording()
            messagebox.showerror("エラー", f"{message}\n録音を停止しました。")
            return
        self.record_poll_id = self.root.after(150, self._poll_recording)

    def _finalize_recording(self):
        """録音ストリーム・ライタースレッド・WAVファイルを後始末する"""
        if self.record_poll_id is not None:
            self.root.after_cancel(self.record_poll_id)
            self.record_poll_id = None
        try:
            if self.record_stream is not None:
                self.record_stream.stop()
                self.record_stream.close()
        except Exception:
            logger.exception("録音ストリームの停止に失敗しました。")
        self.record_stream = None

        if self.record_queue is not None:
            self.record_queue.put(None)
        if self.record_writer_thread is not None:
            self.record_writer_thread.join(timeout=5)
        self.record_writer_thread = None

        try:
            if self.record_wave_file is not None:
                self.record_wave_file.close()
        except Exception:
            logger.exception("録音ファイルのクローズに失敗しました。")
        self.record_wave_file = None

        self.is_recording = False
        self.record_time_label.config(text="00:00")
        self.record_level_bar['value'] = 0
        self._set_recording_ui(False)

    def stop_recording(self, on_close=False):
        """マイク録音を停止する（on_close=True の場合はダイアログを出さず終了処理のみ行う）"""
        if not self.is_recording:
            return
        output_path = self.record_output_path
        self._finalize_recording()
        logger.info(f"録音を保存しました: {output_path}")
        if on_close:
            return
        self.input_file_paths = [output_path]
        display_path = output_path if len(output_path) < 50 else "..." + output_path[-47:]
        self.file_label.config(text=display_path)
        if messagebox.askyesno("録音完了", "録音を保存しました。このまま文字起こしを開始しますか？"):
            self.start_processing()

    def select_input_file(self):
        filetypes = [("音声ファイル", "*.wav;*.mp3;*.m4a;*.mp4")]
        file_paths = filedialog.askopenfilenames(
            title="音声ファイルを選択（複数選択可）",
            filetypes=filetypes,
            initialdir='./'
        )
        if file_paths:
            self.input_file_paths = list(file_paths)
            if len(self.input_file_paths) == 1:
                file_path = self.input_file_paths[0]
                # 長いパスは省略表示
                display_path = file_path if len(file_path) < 50 else "..." + file_path[-47:]
                self.file_label.config(text=display_path)
            else:
                first_name = os.path.basename(self.input_file_paths[0])
                self.file_label.config(
                    text=f"{len(self.input_file_paths)} 件選択: {first_name} ほか"
                )
    
    def select_output_folder(self):
        folder_path = filedialog.askdirectory(title="出力するフォルダを選択", initialdir='./')
        if folder_path:
            self.output_folder_path = folder_path
            display_path = folder_path if len(folder_path) < 50 else "..." + folder_path[-47:]
            self.folder_label.config(text=display_path)
    
    def populate_hotwords_listbox(self):
        """リストボックスに登録済み単語を表示"""
        self.hotwords_listbox.delete(0, tk.END)
        for word in self.hotwords_list:
            self.hotwords_listbox.insert(tk.END, word)
        self.update_hotwords_count()

    def update_hotwords_count(self):
        """登録数カウンターを更新"""
        count = len(self.hotwords_list)
        self.hotwords_count_label.config(text=f"{count} / {MAX_HOTWORDS} 件")

    def add_hotword(self):
        """単語を追加"""
        word = self.hotword_entry.get().strip()
        if not word:
            return
        if len(self.hotwords_list) >= MAX_HOTWORDS:
            messagebox.showwarning(
                "上限",
                f"登録できる単語は最大{MAX_HOTWORDS}件です。\n"
                f"不要な単語を削除してから追加してください。"
            )
            return
        if word in self.hotwords_list:
            messagebox.showinfo("情報", f"「{word}」は既に登録されています。")
            return
        self.hotwords_list.append(word)
        save_hotwords(self.hotwords_list)
        self.hotwords_listbox.insert(tk.END, word)
        self.hotword_entry.delete(0, tk.END)
        self.update_hotwords_count()

    def remove_hotword(self):
        """選択した単語を削除"""
        selected = self.hotwords_listbox.curselection()
        if not selected:
            messagebox.showinfo("情報", "削除する単語を選択してください。")
            return
        # 逆順で削除（インデックスずれ防止）
        for idx in reversed(selected):
            word = self.hotwords_listbox.get(idx)
            self.hotwords_list.remove(word)
            self.hotwords_listbox.delete(idx)
        save_hotwords(self.hotwords_list)
        self.update_hotwords_count()

    def get_hotwords_string(self):
        """登録済み単語をhotwordsパラメータ用の文字列に変換"""
        if not self.hotwords_list:
            return None
        return " ".join(self.hotwords_list)

    def on_model_selected(self, event=None):
        """モデル選択変更時の処理"""
        self.selected_model_name = self.model_combo.get()
        self.settings["model_name"] = self.selected_model_name
        save_settings(self.settings)

    def on_output_option_changed(self):
        """出力オプション変更時の処理"""
        self.settings["output_split"] = self.output_split_var.get()
        self.settings["output_txt"] = self.output_txt_var.get()
        self.settings["output_srt"] = self.output_srt_var.get()
        self.settings["output_docx"] = self.output_docx_var.get()
        self.settings["low_quality_mode"] = self.low_quality_var.get()
        self.settings["diarization"] = self.diarization_var.get()
        save_settings(self.settings)

    def update_progress(self, current, total, status_text, detail_text=""):
        """プログレスバーと状態表示を更新"""
        progress_value = (current / total) * 100 if total > 0 else 0
        self.progress_bar['value'] = progress_value
        self.status_label.config(text=status_text)
        self.progress_detail.config(text=detail_text)
        self.root.update_idletasks()
    
    def set_ui_state(self, enabled):
        """UIの有効/無効を切り替え"""
        state = tk.NORMAL if enabled else tk.DISABLED
        self.file_btn.config(state=state)
        self.folder_btn.config(state=state)
        self.run_btn.config(state=state)
        self.add_btn.config(state=state)
        self.remove_btn.config(state=state)
        self.hotword_entry.config(state=state)
        self.output_split_check.config(state=state)
        self.output_txt_check.config(state=state)
        self.output_srt_check.config(state=state)
        self.output_docx_check.config(state=state)
        self.low_quality_check.config(state=state)
        if self.diarization_model_path:
            self.diarization_check.config(state=state)
        if hasattr(self, "model_combo"):
            self.model_combo.config(state="readonly" if enabled else tk.DISABLED)
        self.cancel_btn.config(state=tk.DISABLED if enabled else tk.NORMAL)
        self.record_start_btn.config(state=tk.DISABLED if (not enabled or self.is_recording) else tk.NORMAL)

    def _set_recording_ui(self, recording):
        """録音中は入力ファイル選択・出力フォルダ選択・文字起こし開始を無効化する"""
        lock_state = tk.DISABLED if recording else tk.NORMAL
        self.file_btn.config(state=lock_state)
        self.folder_btn.config(state=lock_state)
        self.run_btn.config(state=lock_state)
        self.record_start_btn.config(state=tk.DISABLED if recording else tk.NORMAL)
        self.record_stop_btn.config(state=tk.NORMAL if recording else tk.DISABLED)

    def start_processing(self):
        if self.is_recording:
            messagebox.showwarning("警告", "録音中は文字起こしを開始できません。録音を停止してください。")
            return
        if not self.input_file_paths:
            messagebox.showwarning("警告", "入力ファイルを選択してください。")
            return
        if not self.output_folder_path:
            messagebox.showwarning("警告", "出力フォルダを選択してください。")
            return

        existing_names = []
        for file_path in self.input_file_paths:
            file_name = os.path.splitext(os.path.basename(file_path))[0]
            output_file = os.path.join(self.output_folder_path, f"{file_name}_output.xlsx")
            if os.path.exists(output_file):
                existing_names.append(f"{file_name}_output.xlsx")
        if existing_names:
            names_text = "\n".join(existing_names)
            if not messagebox.askyesno(
                "確認",
                f"以下の出力ファイルが既に存在します。上書きしますか？\n\n{names_text}"
            ):
                return

        # 処理を別スレッドで実行（UIフリーズ防止）
        self.cancel_event.clear()
        self.is_processing = True
        self.set_ui_state(False)

        thread = threading.Thread(target=self.process_audio)
        thread.daemon = True
        thread.start()

    def process_audio(self):
        file_count = len(self.input_file_paths)
        succeeded = 0
        failed_names = []
        model = None
        diarization_pipeline = None
        want_diarization = bool(self.diarization_var.get() and self.diarization_model_path)
        try:
            for file_index, file_path in enumerate(self.input_file_paths, start=1):
                if self.cancel_event.is_set():
                    raise ProcessingCancelled()
                file_name = os.path.basename(file_path)
                try:
                    if model is None:
                        self.root.after(0, lambda: self.update_progress(
                            0, 100, "モデルを読み込み中...", "初回は時間がかかる場合があります"))
                        # 同梱モデルのローカルパスを直接指定して読み込む（ネットワークに出ない）。
                        model_path = resolve_local_model_path(self.selected_model_name)
                        if model_path is None:
                            raise RuntimeError(
                                "モデルが見つかりません。アプリケーションを再インストールしてください。"
                            )
                        model = WhisperModel(
                            model_path, device="cpu",
                            compute_type="int8", local_files_only=True
                        )
                    if want_diarization and diarization_pipeline is None:
                        # 話者分離パイプラインはバッチ全体で1回だけロードする。
                        # torch / pyannote は重いため、実際に必要になった時点で初めて import する。
                        try:
                            self.root.after(0, lambda: self.update_progress(
                                0, 100, "話者分離モデルを読み込み中...", ""))
                            from pyannote.audio import Pipeline
                            diarization_pipeline = Pipeline.from_pretrained(
                                DIARIZATION_MODEL_REPO, cache_dir=self.diarization_model_path
                            )
                        except Exception:
                            logger.exception("話者分離モデルの読み込みに失敗しました。話者分離なしで続行します。")
                            want_diarization = False
                            diarization_pipeline = None
                    self.transcribe_file(
                        model, file_path, self.output_folder_path,
                        file_index=file_index, file_count=file_count,
                        diarization_pipeline=diarization_pipeline if want_diarization else None
                    )
                    succeeded += 1
                except ProcessingCancelled:
                    raise
                except Exception:
                    logger.exception(f"文字起こし処理に失敗しました: {file_path}")
                    failed_names.append(file_name)

            self.root.after(0, lambda: self.on_process_complete(file_count, succeeded, failed_names))
        except ProcessingCancelled:
            logger.info("ユーザー操作により処理がキャンセルされました。")
            self.root.after(0, lambda: messagebox.showinfo("キャンセル", "処理をキャンセルしました。"))
        finally:
            self.is_processing = False
            self.root.after(0, lambda: self.set_ui_state(True))
            self.root.after(0, lambda: self.update_progress(0, 1, "待機中...", ""))

    def on_process_complete(self, file_count, succeeded, failed_names):
        """処理完了時の処理"""
        self.notify_completion()
        if not failed_names:
            messagebox.showinfo("完了", "処理が完了しました。")
        else:
            names_text = "\n".join(failed_names)
            messagebox.showwarning(
                "完了",
                f"{file_count}件中{succeeded}件成功。\n失敗: {names_text}\n\n"
                f"詳細はログファイル (logs フォルダ) を確認してください。"
            )
        # 出力フォルダを開く
        if self.output_folder_path:
            os.startfile(self.output_folder_path)

    def notify_completion(self):
        """処理完了を音とタスクバー点滅で通知する"""
        try:
            winsound.MessageBeep(winsound.MB_ICONASTERISK)
            FLASHW_ALL = 0x00000003
            FLASHW_TIMERNOFG = 0x0000000C

            class FLASHWINFO(ctypes.Structure):
                _fields_ = [
                    ("cbSize", ctypes.c_uint),
                    ("hwnd", ctypes.c_void_p),
                    ("dwFlags", ctypes.c_uint),
                    ("uCount", ctypes.c_uint),
                    ("dwTimeout", ctypes.c_uint),
                ]

            hwnd = ctypes.windll.user32.GetParent(self.root.winfo_id())
            info = FLASHWINFO(
                ctypes.sizeof(FLASHWINFO), hwnd,
                FLASHW_ALL | FLASHW_TIMERNOFG, 5, 0
            )
            ctypes.windll.user32.FlashWindowEx(ctypes.byref(info))
        except Exception:
            pass

    @staticmethod
    def format_time(seconds):
        """秒数を HH:MM:SS 形式の文字列に変換"""
        sec = int(seconds)
        return f"{sec // 3600:02d}:{(sec % 3600) // 60:02d}:{sec % 60:02d}"

    @staticmethod
    def format_srt_time(seconds):
        """秒数を SRT 用の HH:MM:SS,mmm 形式の文字列に変換"""
        total_ms = int(round(seconds * 1000))
        ms = total_ms % 1000
        total_sec = total_ms // 1000
        h = total_sec // 3600
        m = (total_sec % 3600) // 60
        s = total_sec % 60
        return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"

    def diarize_and_assign(self, diarization_pipeline, file_path, raw_segments, prefix):
        """話者分離を実行し、raw_segments と同じ順序で話者ラベル（話者A等）のリストを返す"""
        import torch

        audio = decode_audio_pyav_16k_mono(file_path)
        waveform = torch.from_numpy(audio).unsqueeze(0)
        audio_input = {"waveform": waveform, "sample_rate": DIARIZATION_SAMPLE_RATE}

        def hook(step_name, step_artifact, file=None, total=None, completed=None):
            if self.cancel_event.is_set():
                raise ProcessingCancelled()
            if total and completed is not None:
                detail_text = f"話者分離中: {step_name} {completed}/{total}"
            else:
                detail_text = f"話者分離中: {step_name}"
            self.root.after(0, lambda dt=detail_text: self.update_progress(
                96, 100, f"{prefix}話者分離中...", dt))

        diarization = diarization_pipeline(audio_input, hook=hook)
        annotation = getattr(diarization, "speaker_diarization", diarization)
        turns = [
            {"start": turn.start, "end": turn.end, "speaker": speaker}
            for turn, _, speaker in annotation.itertracks(yield_label=True)
        ]

        # 話者ラベルは登場順（区間の開始時刻順）に話者A、話者B…へ変換する
        speaker_order = []
        for turn in sorted(turns, key=lambda t: t["start"]):
            if turn["speaker"] not in speaker_order:
                speaker_order.append(turn["speaker"])
        label_map = {}
        for i, raw_speaker in enumerate(speaker_order):
            label_map[raw_speaker] = f"話者{chr(ord('A') + i)}" if i < 26 else f"話者{i + 1}"

        labels = []
        for seg_start, seg_end, _text in raw_segments:
            raw_speaker = assign_speaker_to_segment(seg_start, seg_end, turns)
            labels.append(label_map.get(raw_speaker, DIARIZATION_UNKNOWN_SPEAKER))
        return labels

    def transcribe_file(self, model, file_path, output_folder, file_index=1, file_count=1,
                         diarization_pipeline=None):
        file_extension = os.path.splitext(file_path)[1].lower()
        if file_extension not in (".wav", ".mp3", ".m4a", ".mp4"):
            raise ValueError("サポートされていない音声形式です。")
        file_name = os.path.splitext(os.path.basename(file_path))[0]

        prefix = f"[{file_index}/{file_count}] {os.path.basename(file_path)}: " if file_count > 1 else ""

        os.makedirs(output_folder, exist_ok=True)

        start_time_all = datetime.datetime.now()
        logger.info(f"処理開始: {file_path}")

        transcribe_params = dict(
            beam_size=5,
            language='ja',
            vad_filter=True,
            initial_prompt=INITIAL_PROMPT,
            # 音質の悪い音源で一度誤認識（幻覚）が起きると、直前の出力を文脈として
            # 引き継ぐ仕組みにより残り全体へ連鎖するため無効化する。
            # v1.3の「1分ごとの独立認識」が持っていたリセット効果に相当。
            # temperature は既定のフォールバック（失敗時に温度を上げて再試行）を使う。
            condition_on_previous_text=False,
        )
        # hotwords は30秒ごとの認識窓すべてのプロンプトに注入されるため、
        # 句読点誘導文を載せることで全区間に句読点が付くようにする
        # （initial_prompt は先頭の窓にしか効かない）。ユーザー登録単語も併記する。
        hotwords_str = self.get_hotwords_string()
        transcribe_params["hotwords"] = (
            f"{INITIAL_PROMPT} {hotwords_str}" if hotwords_str else INITIAL_PROMPT
        )

        # 低品質音源モード: ノイズ抑制+音量正規化した波形を直接認識にかける。
        # 通常音源ではわずかに悪化しうるため既定OFF（雑音がひどい音源の救済用）。
        audio_input = file_path
        if self.low_quality_var.get():
            self.root.after(0, lambda: self.update_progress(
                3, 100, f"{prefix}音声を前処理中（ノイズ抑制）..."))
            logger.info("低品質音源モード: ノイズ抑制前処理を実行")
            audio_input = preprocess_low_quality(
                decode_audio(file_path, sampling_rate=16000), self.cancel_event)
            transcribe_params["vad_parameters"] = dict(threshold=0.25)

        # ファイル全体を1回で文字起こしし、タイムスタンプで1分単位にまとめる
        # （物理分割しないため文の途中で切れず、再エンコードによる劣化もない）
        self.root.after(0, lambda: self.update_progress(5, 100, f"{prefix}文字起こし中...", ""))
        segments, info = model.transcribe(audio_input, **transcribe_params)
        duration = max(info.duration or 0, 1.0)
        logger.info(f"音声長: {duration:.1f}秒")

        split_interval = 60
        buckets = {}
        bucket_logprobs = {}
        raw_segments = []
        total_chars = 0
        segment_loop_start = datetime.datetime.now()
        for segment in segments:
            if self.cancel_event.is_set():
                raise ProcessingCancelled()
            idx = int(segment.start // split_interval)
            text = str(segment.text).strip()
            buckets.setdefault(idx, []).append(text)
            raw_segments.append((segment.start, segment.end, text))
            total_chars += len(text)
            avg_logprob = segment.avg_logprob
            if idx not in bucket_logprobs or avg_logprob < bucket_logprobs[idx]:
                bucket_logprobs[idx] = avg_logprob

            elapsed = (datetime.datetime.now() - segment_loop_start).total_seconds()
            detail_text = ""
            if elapsed >= 1.0 and segment.end > 0:
                speed = segment.end / elapsed
                if speed > 0:
                    remaining = max(duration - segment.end, 0) / speed
                    if remaining >= 60:
                        detail_text = f"残り約{int(remaining // 60)}分"
                    else:
                        detail_text = f"残り約{int(remaining)}秒"

            progress = 5 + min(segment.end / duration, 1.0) * 93
            status_text = f"{prefix}文字起こし中... ({self.format_time(segment.end)} / {self.format_time(duration)})"
            self.root.after(0, lambda p=progress, st=status_text, dt=detail_text:
                            self.update_progress(p, 100, st, dt))

        low_conf_rows = {idx for idx, v in bucket_logprobs.items() if v < LOW_CONFIDENCE_LOGPROB}
        num_rows = max(buckets.keys()) + 1 if buckets else 1
        end_time_all = datetime.datetime.now()
        logger.info(f"文字起こし完了 ({total_chars}文字) 処理時間: {end_time_all - start_time_all}")

        # 話者分離: Whisperの認識結果には一切手を加えず、独立に実行した結果を後合成する。
        row_speakers = {}
        segment_speaker_labels = None
        diarization_succeeded = False
        if diarization_pipeline is not None:
            self.root.after(0, lambda: self.update_progress(96, 100, f"{prefix}話者分離中..."))
            try:
                speaker_labels = self.diarize_and_assign(
                    diarization_pipeline, file_path, raw_segments, prefix)
                for (seg_start, _seg_end, _text), label in zip(raw_segments, speaker_labels):
                    idx = int(seg_start // split_interval)
                    bucket_speakers = row_speakers.setdefault(idx, [])
                    if label not in bucket_speakers:
                        bucket_speakers.append(label)
                # 正規の話者が1人でもいるバケットからは「不明」を除く
                for idx in list(row_speakers.keys()):
                    real = [sp for sp in row_speakers[idx] if sp != DIARIZATION_UNKNOWN_SPEAKER]
                    if real:
                        row_speakers[idx] = real
                diarization_succeeded = True
                segment_speaker_labels = speaker_labels
                logger.info(f"話者分離完了: {file_path}")
            except ProcessingCancelled:
                raise
            except Exception:
                logger.exception(f"話者分離に失敗しました: {file_path}")
                row_speakers = {}
                self.root.after(0, lambda: self.update_progress(
                    96, 100, f"{prefix}話者分離に失敗しました。話者なしで続行します。"))
        show_speaker_column = diarization_succeeded

        # 再生用の分割音声を専用フォルダに出力（Excelの各行から該当区間へ頭出しできるようにする）
        row_links = {}
        if self.output_split_var.get():
            self.root.after(0, lambda: self.update_progress(
                95, 100, f"{prefix}再生用の分割音声を出力中..."))
            split_dir = os.path.join(output_folder, f"{file_name}_分割音声")
            os.makedirs(split_dir, exist_ok=True)
            for idx in range(num_rows):
                if self.cancel_event.is_set():
                    raise ProcessingCancelled()
                start_sec = idx * split_interval
                end_sec = min((idx + 1) * split_interval, duration)
                split_path = os.path.join(split_dir, f"{file_name}_{idx}{file_extension}")
                try:
                    export_audio_segment(file_path, split_path, start_sec, end_sec)
                    row_links[idx] = split_path
                except Exception:
                    logger.exception(f"分割音声の出力に失敗しました: {split_path}")
            logger.info(f"分割音声 {len(row_links)}/{num_rows} 件を {split_dir} に出力しました。")

        # 進捗更新: Excel出力
        self.root.after(0, lambda: self.update_progress(98, 100, f"{prefix}Excelファイルを出力中..."))

        output_file = os.path.join(output_folder, f"{file_name}_output.xlsx")
        workbook = Workbook()
        sheet = workbook.active
        if show_speaker_column:
            sheet.append(['No', '時間帯', '話者', '音声ファイル', '変換結果'])
            link_col, text_col = 4, 5
        else:
            sheet.append(['No', '時間帯', '音声ファイル', '変換結果'])
            link_col, text_col = 3, 4
        low_conf_fill = PatternFill(fill_type="solid", start_color="FFF9C4")
        for idx in range(num_rows):
            start_sec = idx * split_interval
            end_sec = min((idx + 1) * split_interval, duration)
            time_label = f"{self.format_time(start_sec)} - {self.format_time(end_sec)}"
            link_path = row_links.get(idx)
            link_name = os.path.basename(link_path) if link_path else ""
            text_value = '\n'.join(buckets.get(idx, []))
            if show_speaker_column:
                speaker_label = '、'.join(row_speakers.get(idx, []))
                sheet.append([str(idx), time_label, speaker_label, link_name, text_value])
            else:
                sheet.append([str(idx), time_label, link_name, text_value])
            # 音声ファイルセルから該当区間の分割音声を開けるようにリンクを付与
            if link_path:
                sheet.cell(row=idx + 2, column=link_col).hyperlink = link_path
            if idx in low_conf_rows:
                sheet.cell(row=idx + 2, column=text_col).fill = low_conf_fill
        sheet.column_dimensions['A'].width = 6
        sheet.column_dimensions['B'].width = 22
        if show_speaker_column:
            sheet.column_dimensions['C'].width = 16
            sheet.column_dimensions['D'].width = 28
            sheet.column_dimensions['E'].width = 100
        else:
            sheet.column_dimensions['C'].width = 28
            sheet.column_dimensions['D'].width = 100
        for row in sheet.iter_rows(min_row=2, min_col=text_col, max_col=text_col):
            row[0].alignment = Alignment(wrap_text=True, vertical='top')

        try:
            workbook.save(output_file)
        except PermissionError:
            raise RuntimeError(
                f"Excelファイルを保存できません。\n{output_file}\n"
                f"このファイルを開いている場合は閉じてから再実行してください。"
            )
        logger.info(f"Excelファイル {output_file} を保存しました。")

        if diarization_succeeded and segment_speaker_labels is not None:
            self.write_speakers_output(output_folder, file_name, raw_segments, segment_speaker_labels)

        if self.output_txt_var.get():
            self.write_txt_output(output_folder, file_name, buckets, num_rows, split_interval, duration)
        if self.output_srt_var.get():
            self.write_srt_output(output_folder, file_name, raw_segments)
        if self.output_docx_var.get():
            self.write_docx_output(output_folder, file_name, buckets, num_rows, split_interval, duration, low_conf_rows,
                                    row_speakers=row_speakers, show_speaker_column=show_speaker_column)

        # 進捗更新: 完了
        self.root.after(0, lambda: self.update_progress(100, 100, f"{prefix}完了！",
                       f"総処理時間: {end_time_all - start_time_all}"))

    def write_speakers_output(self, output_folder, file_name, raw_segments, speaker_labels):
        """話者分離結果を発言単位（連続する同一話者を1行に結合）でExcel出力する"""
        rows = []
        for (seg_start, seg_end, text), speaker in zip(raw_segments, speaker_labels):
            if rows and rows[-1]['speaker'] == speaker:
                rows[-1]['end'] = seg_end
                rows[-1]['texts'].append(text)
            else:
                rows.append({'start': seg_start, 'end': seg_end, 'speaker': speaker, 'texts': [text]})

        output_file = os.path.join(output_folder, f"{file_name}_speakers.xlsx")
        workbook = Workbook()
        sheet = workbook.active
        sheet.append(['No', '時間帯', '話者', '発言内容'])
        for i, row in enumerate(rows, start=1):
            time_label = f"{self.format_time(row['start'])} - {self.format_time(row['end'])}"
            text_value = '\n'.join(row['texts'])
            sheet.append([i, time_label, row['speaker'], text_value])
        sheet.column_dimensions['A'].width = 6
        sheet.column_dimensions['B'].width = 22
        sheet.column_dimensions['C'].width = 16
        sheet.column_dimensions['D'].width = 100
        for row_cells in sheet.iter_rows(min_row=2, min_col=4, max_col=4):
            row_cells[0].alignment = Alignment(wrap_text=True, vertical='top')

        try:
            workbook.save(output_file)
        except PermissionError:
            raise RuntimeError(
                f"Excelファイルを保存できません。\n{output_file}\n"
                f"このファイルを開いている場合は閉じてから再実行してください。"
            )
        logger.info(f"話者別Excelファイル {output_file} を保存しました。")

    def write_txt_output(self, output_folder, file_name, buckets, num_rows, split_interval, duration):
        """1分ブロック単位のテキストファイルを出力"""
        output_file = os.path.join(output_folder, f"{file_name}_output.txt")
        lines = []
        for idx in range(num_rows):
            start_sec = idx * split_interval
            end_sec = min((idx + 1) * split_interval, duration)
            lines.append(f"[{self.format_time(start_sec)} - {self.format_time(end_sec)}]")
            lines.append('\n'.join(buckets.get(idx, [])))
            lines.append("")
        with open(output_file, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        logger.info(f"テキストファイル {output_file} を保存しました。")

    def write_srt_output(self, output_folder, file_name, raw_segments):
        """セグメント単位のSRT字幕ファイルを出力"""
        output_file = os.path.join(output_folder, f"{file_name}_output.srt")
        lines = []
        for i, (start, end, text) in enumerate(raw_segments, start=1):
            lines.append(str(i))
            lines.append(f"{self.format_srt_time(start)} --> {self.format_srt_time(end)}")
            lines.append(text)
            lines.append("")
        with open(output_file, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        logger.info(f"字幕ファイル {output_file} を保存しました。")

    def write_docx_output(self, output_folder, file_name, buckets, num_rows, split_interval, duration, low_conf_rows,
                           row_speakers=None, show_speaker_column=False):
        """1分ブロック単位のWord文書を出力"""
        row_speakers = row_speakers or {}
        output_file = os.path.join(output_folder, f"{file_name}_output.docx")
        document = docx.Document()
        document.add_heading(f"文字起こし結果: {file_name}", level=1)
        document.add_paragraph(
            f"元ファイル: {file_name} / 総時間: {self.format_time(duration)} / "
            f"作成日時: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}"
        )
        for idx in range(num_rows):
            start_sec = idx * split_interval
            end_sec = min((idx + 1) * split_interval, duration)
            heading_text = f"[{self.format_time(start_sec)} - {self.format_time(end_sec)}]"
            if show_speaker_column and row_speakers.get(idx):
                heading_text += f" 話者: {'、'.join(row_speakers[idx])}"
            if idx in low_conf_rows:
                heading_text += " ※要確認"
            heading_para = document.add_paragraph()
            heading_run = heading_para.add_run(heading_text)
            heading_run.bold = True
            document.add_paragraph('\n'.join(buckets.get(idx, [])))
            document.add_paragraph("")

        try:
            document.save(output_file)
        except PermissionError:
            raise RuntimeError(
                f"Wordファイルを保存できません。\n{output_file}\n"
                f"このファイルを開いている場合は閉じてから再実行してください。"
            )
        logger.info(f"Wordファイル {output_file} を保存しました。")

    def copy_support_info(self):
        """サポート情報をまとめてクリップボードにコピーする"""
        is_frozen = getattr(sys, 'frozen', False)
        lines = [
            APP_TITLE,
            f"実行形態: {'PyInstaller' if is_frozen else '開発'}",
            f"OS: {platform.platform()}",
            f"使用モデル: {self.selected_model_name}",
            f"検出モデル一覧: {', '.join(self.available_models) if self.available_models else 'なし'}",
            f"単語登録数: {len(self.hotwords_list)} / {MAX_HOTWORDS}",
            f"出力オプション: 分割音声={self.output_split_var.get()}, "
            f"txt={self.output_txt_var.get()}, srt={self.output_srt_var.get()}, "
            f"docx={self.output_docx_var.get()}",
            f"話者分離: {self.diarization_var.get()} (モデル導入={bool(self.diarization_model_path)})",
        ]

        log_dir = os.path.join(get_app_dir(), "logs")
        log_file = os.path.join(log_dir, f"app-{datetime.datetime.now().strftime('%Y%m%d')}.log")
        if os.path.exists(log_file):
            try:
                with open(log_file, "r", encoding="utf-8") as f:
                    log_lines = f.readlines()
                lines.append("")
                lines.append("--- ログ末尾40行 ---")
                lines.extend(line.rstrip("\n") for line in log_lines[-40:])
            except OSError:
                lines.append("(ログ読込不可)")

        info_text = "\n".join(lines)
        self.root.clipboard_clear()
        self.root.clipboard_append(info_text)
        messagebox.showinfo("サポート情報", "サポート情報をコピーしました。メール等に貼り付けてご利用ください。")

    def run(self):
        self.root.mainloop()


def main():
    if "--selftest" in sys.argv:
        try:
            setup_logging()
        except Exception:
            pass
        sys.exit(run_selftest())
    try:
        setup_logging()
    except Exception:
        pass
    app = AudioTranscriptionApp()
    app.run()


if __name__ == "__main__":
    main()
